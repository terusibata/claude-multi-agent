# -*- coding: utf-8 -*-
"""
Wordファイル用ツール

AIエージェントがWordファイルを理解するための軽量ツール。
3つの機能を提供:
1. get_document_info: ドキュメント構造と基本情報を取得
2. get_document_content: 指定範囲の内容を取得（段落範囲指定可能）
3. search_document: ドキュメント全体からキーワード検索
"""

from __future__ import annotations

import io
import re
import unicodedata
from typing import TYPE_CHECKING, Any, Optional, TypedDict

import structlog

if TYPE_CHECKING:
    from app.services.workspace_service import WorkspaceService

logger = structlog.get_logger(__name__)


# =============================================================================
# Type Definitions
# =============================================================================

class HeadingInfo(TypedDict):
    """見出し情報の型定義"""
    level: int
    text: str
    para_index: int
    char_count: int  # この見出しセクションの文字数


class TableInfo(TypedDict):
    """表情報の型定義"""
    index: int
    rows: int
    cols: int
    near_para: int  # 近くの段落番号（おおよその位置）


class DocumentInfo(TypedDict):
    """ドキュメント情報の型定義"""
    filename: str
    total_paragraphs: int
    total_characters: int
    tables_count: int
    headings: list[HeadingInfo]
    tables: list[TableInfo]


class ContentResult(TypedDict):
    """コンテンツ取得結果の型定義"""
    filename: str
    section_title: str | None
    start_paragraph: int
    end_paragraph: int
    total_paragraphs: int
    returned_paragraphs: int
    has_more: bool
    content: str
    tables_in_range: int


class SearchHit(TypedDict):
    """検索ヒットの型定義"""
    location_type: str  # "paragraph" | "table" | "heading"
    para_index: int | None
    table_index: int | None
    heading_level: int | None
    text: str
    context: str


class SearchResult(TypedDict):
    """検索結果の型定義"""
    query: str
    total_hits: int
    hits: list[SearchHit]


# =============================================================================
# Constants
# =============================================================================

DEFAULT_MAX_PARAGRAPHS = 50  # デフォルトの最大取得段落数


# =============================================================================
# Internal Utilities
# =============================================================================

def _normalize_text(text: str) -> str:
    """テキストの正規化（Unicode NFC、制御文字除去）"""
    if not text:
        return ""

    text = unicodedata.normalize("NFC", text)

    result = []
    for ch in text:
        code = ord(ch)
        if ch in ('\n', '\r', '\t'):
            result.append(ch)
        elif code < 0x20 or code == 0x7F or (0x80 <= code <= 0x9F):
            continue
        elif code in (0x200B, 0x200C, 0x200D, 0x2060, 0xFEFF):
            continue
        else:
            result.append(ch)

    return "".join(result)


def _get_heading_level(para) -> int | None:
    """段落が見出しなら、そのレベルを返す"""
    style_name = para.style.name if para.style else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.replace("Heading ", "").replace("Heading", "1"))
        except ValueError:
            return 1
    return None


def _load_document_from_bytes(content: bytes):
    """バイトデータからドキュメントを読み込む"""
    from docx import Document
    return Document(io.BytesIO(content))


def _check_old_format(file_path: str) -> dict[str, Any] | None:
    """古いOffice形式（.doc）のチェック"""
    if file_path.lower().endswith(".doc"):
        return {
            "content": [{
                "type": "text",
                "text": (
                    f"エラー: '{file_path}' は古いWord形式（.doc）です。\n\n"
                    "このツールは .docx（Office Open XML）形式のみ対応しています。\n"
                    ".doc（バイナリ形式）ファイルは python-docx では読み取れません。\n\n"
                    "対処方法:\n"
                    "1. Microsoft Word で .docx 形式に変換して再アップロード\n"
                    "2. LibreOffice で .docx 形式に変換して再アップロード\n"
                    "3. オンライン変換ツールを使用"
                ),
            }],
            "is_error": True,
        }
    return None


def _create_context_snippet(text: str, match_start: int, match_end: int, context_chars: int = 40) -> str:
    """検索マッチの前後コンテキストを生成"""
    start = max(0, match_start - context_chars)
    end = min(len(text), match_end + context_chars)

    prefix = "..." if start > 0 else ""
    suffix = "..." if end < len(text) else ""

    before = text[start:match_start]
    match = text[match_start:match_end]
    after = text[match_end:end]

    return f"{prefix}{before}[{match}]{after}{suffix}"


# =============================================================================
# Core Functions
# =============================================================================

def get_document_info(content: bytes, filename: str) -> DocumentInfo:
    """
    Wordドキュメントの構造情報を取得する。

    Args:
        content: Wordファイルのバイトデータ
        filename: ファイル名

    Returns:
        DocumentInfo: ドキュメント情報
    """
    doc = _load_document_from_bytes(content)

    paragraphs = list(doc.paragraphs)
    total_paragraphs = len(paragraphs)
    total_characters = 0

    # 見出し情報を収集
    headings: list[HeadingInfo] = []
    current_heading_start = 0
    current_heading_chars = 0

    for i, para in enumerate(paragraphs):
        para_text = _normalize_text(para.text)
        char_count = len(para_text)
        total_characters += char_count

        level = _get_heading_level(para)
        if level is not None:
            # 前の見出しの文字数を確定
            if headings:
                headings[-1]["char_count"] = current_heading_chars

            headings.append(HeadingInfo(
                level=level,
                text=para_text.strip()[:100],  # 見出しテキストは100文字まで
                para_index=i + 1,
                char_count=0,  # 後で更新
            ))
            current_heading_chars = char_count
        else:
            current_heading_chars += char_count

    # 最後の見出しの文字数を確定
    if headings:
        headings[-1]["char_count"] = current_heading_chars

    # 表情報を収集
    tables: list[TableInfo] = []
    for idx, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0

        # 表の位置を推定（簡易的に）
        tables.append(TableInfo(
            index=idx + 1,
            rows=rows,
            cols=cols,
            near_para=min(total_paragraphs, (idx + 1) * (total_paragraphs // (len(doc.tables) + 1))),
        ))

    return DocumentInfo(
        filename=filename,
        total_paragraphs=total_paragraphs,
        total_characters=total_characters,
        tables_count=len(doc.tables),
        headings=headings,
        tables=tables,
    )


def get_document_content(
    content: bytes,
    *,
    heading: str | None = None,
    start_paragraph: int | None = None,
    end_paragraph: int | None = None,
    max_paragraphs: int = DEFAULT_MAX_PARAGRAPHS,
    include_tables: bool = True,
) -> ContentResult:
    """
    指定範囲のドキュメント内容を取得する。

    Args:
        content: Wordファイルのバイトデータ
        heading: 見出しテキスト（部分一致検索）
        start_paragraph: 開始段落番号（1始まり）
        end_paragraph: 終了段落番号（1始まり）
        max_paragraphs: 最大取得段落数（デフォルト: 50）
        include_tables: 表を含めるか（デフォルト: True）

    Returns:
        ContentResult: 取得結果
    """
    doc = _load_document_from_bytes(content)
    paragraphs = list(doc.paragraphs)
    total_paragraphs = len(paragraphs)

    section_title = None
    actual_start = 1
    actual_end = total_paragraphs

    if heading:
        # 見出しで検索
        found = False
        for i, para in enumerate(paragraphs):
            level = _get_heading_level(para)
            if level is not None and heading.lower() in para.text.lower():
                found = True
                actual_start = i + 1
                section_title = para.text.strip()[:100]
                current_level = level

                # 次の同レベル以上の見出しまで
                for j in range(i + 1, total_paragraphs):
                    next_level = _get_heading_level(paragraphs[j])
                    if next_level is not None and next_level <= current_level:
                        actual_end = j
                        break
                else:
                    actual_end = total_paragraphs
                break

        if not found:
            raise ValueError(f"見出し '{heading}' が見つかりません。get_document_info でシート一覧を確認してください。")
    else:
        # 段落番号で指定
        if start_paragraph is not None:
            actual_start = max(1, start_paragraph)

        if end_paragraph is not None:
            actual_end = min(total_paragraphs, end_paragraph)
        else:
            actual_end = min(actual_start + max_paragraphs - 1, total_paragraphs)

    # 最大段落数の制限
    if actual_end - actual_start + 1 > max_paragraphs:
        actual_end = actual_start + max_paragraphs - 1

    # コンテンツを抽出
    content_lines: list[str] = []
    tables_in_range = 0

    for i in range(actual_start - 1, min(actual_end, total_paragraphs)):
        para = paragraphs[i]
        para_text = _normalize_text(para.text).strip()
        if para_text:
            content_lines.append(para_text)
            content_lines.append("")

    # 表の処理
    if include_tables and doc.tables:
        tables_in_range = len(doc.tables)  # 簡易的に全表をカウント
        if tables_in_range > 0:
            content_lines.append("")
            content_lines.append(f"[ドキュメントには {tables_in_range} 個の表が含まれています]")

    has_more = actual_end < total_paragraphs

    return ContentResult(
        filename="",  # ハンドラーで設定
        section_title=section_title,
        start_paragraph=actual_start,
        end_paragraph=actual_end,
        total_paragraphs=total_paragraphs,
        returned_paragraphs=actual_end - actual_start + 1,
        has_more=has_more,
        content="\n".join(content_lines),
        tables_in_range=tables_in_range,
    )


def search_document(
    content: bytes,
    query: str,
    *,
    case_sensitive: bool = False,
    max_hits: int = 50,
    include_tables: bool = True,
) -> SearchResult:
    """
    ドキュメント全体からキーワード検索を行う。

    Args:
        content: Wordファイルのバイトデータ
        query: 検索キーワード
        case_sensitive: 大文字小文字を区別するか（デフォルト: False）
        max_hits: 最大ヒット数（デフォルト: 50）
        include_tables: 表も検索対象に含めるか（デフォルト: True）

    Returns:
        SearchResult: 検索結果
    """
    doc = _load_document_from_bytes(content)

    hits: list[SearchHit] = []

    # 検索パターンを準備
    flags = 0 if case_sensitive else re.IGNORECASE
    try:
        pattern = re.compile(re.escape(query), flags)
    except re.error:
        raise ValueError(f"無効な検索クエリ: {query}")

    # 段落を検索
    for i, para in enumerate(doc.paragraphs):
        if len(hits) >= max_hits:
            break

        para_text = _normalize_text(para.text)
        if not para_text:
            continue

        match = pattern.search(para_text)
        if match:
            level = _get_heading_level(para)

            hits.append(SearchHit(
                location_type="heading" if level else "paragraph",
                para_index=i + 1,
                table_index=None,
                heading_level=level,
                text=para_text[:200] if len(para_text) > 200 else para_text,
                context=_create_context_snippet(para_text, match.start(), match.end()),
            ))

    # 表を検索
    if include_tables:
        for table_idx, table in enumerate(doc.tables):
            if len(hits) >= max_hits:
                break

            for row_idx, row in enumerate(table.rows):
                if len(hits) >= max_hits:
                    break

                row_texts = []
                for cell in row.cells:
                    cell_text = _normalize_text(cell.text).strip()
                    row_texts.append(cell_text)

                    match = pattern.search(cell_text)
                    if match and len(hits) < max_hits:
                        # 行全体をコンテキストとして
                        row_context = " | ".join(row_texts)

                        hits.append(SearchHit(
                            location_type="table",
                            para_index=None,
                            table_index=table_idx + 1,
                            heading_level=None,
                            text=cell_text[:200] if len(cell_text) > 200 else cell_text,
                            context=f"表{table_idx + 1}, 行{row_idx + 1}: {row_context[:150]}",
                        ))
                        break  # 1行につき1ヒットまで

    return SearchResult(
        query=query,
        total_hits=len(hits),
        hits=hits,
    )


# =============================================================================
# Tool Handlers
# =============================================================================

async def get_document_info_handler(
    workspace_service: "WorkspaceService",
    tenant_id: str,
    conversation_id: str,
    args: dict[str, Any],
) -> dict[str, Any]:
    """
    Wordドキュメントの構造情報を取得するハンドラー

    Args:
        args:
            file_path: ファイルパス
    """
    file_path = args.get("file_path", "")

    old_format_error = _check_old_format(file_path)
    if old_format_error:
        return old_format_error

    try:
        from docx import Document  # noqa: F401
    except ImportError:
        return {
            "content": [{"type": "text", "text": "エラー: python-docxライブラリがインストールされていません。"}],
            "is_error": True,
        }

    try:
        content, filename, _ = await workspace_service.download_file(
            tenant_id, conversation_id, file_path
        )

        info = get_document_info(content, filename)

        result_lines = [
            f"# Word文書情報: {info['filename']}",
            f"総段落数: {info['total_paragraphs']}",
            f"総文字数: {info['total_characters']:,}",
            f"表の数: {info['tables_count']}",
            "",
        ]

        # 見出し構造
        if info['headings']:
            result_lines.append("## 見出し構造")
            for h in info['headings'][:50]:
                indent = "  " * (h['level'] - 1)
                result_lines.append(f"{indent}- {h['text']} (段落{h['para_index']}, 約{h['char_count']}文字)")
            if len(info['headings']) > 50:
                result_lines.append(f"... 他 {len(info['headings']) - 50} 見出し")
            result_lines.append("")
        else:
            result_lines.append("## 見出し構造")
            result_lines.append("見出しは定義されていません。")
            result_lines.append("")

        # 表の概要
        if info['tables']:
            result_lines.append("## 表の概要")
            for t in info['tables'][:10]:
                result_lines.append(f"- 表{t['index']}: {t['rows']}行 x {t['cols']}列")
            if len(info['tables']) > 10:
                result_lines.append(f"... 他 {len(info['tables']) - 10} 表")
            result_lines.append("")

        result_lines.append("---")
        result_lines.append("データ取得: `get_document_content` を使用")
        result_lines.append("検索: `search_document` を使用")

        return {
            "content": [{"type": "text", "text": "\n".join(result_lines)}],
        }
    except FileNotFoundError:
        return {
            "content": [{"type": "text", "text": f"ファイルが見つかりません: {file_path}"}],
            "is_error": True,
        }
    except Exception as e:
        logger.error("Word情報取得エラー", error=str(e), file_path=file_path)
        return {
            "content": [{"type": "text", "text": f"読み込みエラー: {str(e)}"}],
            "is_error": True,
        }


async def get_document_content_handler(
    workspace_service: "WorkspaceService",
    tenant_id: str,
    conversation_id: str,
    args: dict[str, Any],
) -> dict[str, Any]:
    """
    Wordドキュメントの内容を取得するハンドラー

    Args:
        args:
            file_path: ファイルパス
            heading: 見出しテキスト（部分一致）
            start_paragraph: 開始段落番号（1始まり）
            end_paragraph: 終了段落番号
            max_paragraphs: 最大取得段落数（デフォルト: 50）
            include_tables: 表を含めるか（デフォルト: true）
    """
    file_path = args.get("file_path", "")
    heading = args.get("heading")
    start_paragraph = args.get("start_paragraph")
    end_paragraph = args.get("end_paragraph")
    max_paragraphs = args.get("max_paragraphs", DEFAULT_MAX_PARAGRAPHS)
    include_tables = args.get("include_tables", True)

    old_format_error = _check_old_format(file_path)
    if old_format_error:
        return old_format_error

    try:
        from docx import Document  # noqa: F401
    except ImportError:
        return {
            "content": [{"type": "text", "text": "エラー: python-docxライブラリがインストールされていません。"}],
            "is_error": True,
        }

    try:
        content, filename, _ = await workspace_service.download_file(
            tenant_id, conversation_id, file_path
        )

        result = get_document_content(
            content,
            heading=heading,
            start_paragraph=start_paragraph,
            end_paragraph=end_paragraph,
            max_paragraphs=max_paragraphs,
            include_tables=include_tables,
        )

        result_lines = [
            f"# {filename}",
        ]

        if result['section_title']:
            result_lines.append(f"セクション: {result['section_title']}")

        result_lines.append(f"段落範囲: {result['start_paragraph']}-{result['end_paragraph']} (全{result['total_paragraphs']}段落)")
        result_lines.append(f"取得: {result['returned_paragraphs']}段落")
        result_lines.append("")
        result_lines.append(result['content'])

        if result['has_more']:
            result_lines.append("")
            result_lines.append("---")
            result_lines.append("まだ続きがあります。次を取得するには:")
            result_lines.append(f"`start_paragraph={result['end_paragraph'] + 1}` を指定してください。")

        return {
            "content": [{"type": "text", "text": "\n".join(result_lines)}],
        }
    except FileNotFoundError:
        return {
            "content": [{"type": "text", "text": f"ファイルが見つかりません: {file_path}"}],
            "is_error": True,
        }
    except ValueError as e:
        return {
            "content": [{"type": "text", "text": str(e)}],
            "is_error": True,
        }
    except Exception as e:
        logger.error("Wordコンテンツ取得エラー", error=str(e), file_path=file_path)
        return {
            "content": [{"type": "text", "text": f"読み込みエラー: {str(e)}"}],
            "is_error": True,
        }


async def search_document_handler(
    workspace_service: "WorkspaceService",
    tenant_id: str,
    conversation_id: str,
    args: dict[str, Any],
) -> dict[str, Any]:
    """
    Wordドキュメント全体からキーワード検索を行うハンドラー

    Args:
        args:
            file_path: ファイルパス
            query: 検索キーワード
            case_sensitive: 大文字小文字を区別するか（デフォルト: false）
            max_hits: 最大ヒット数（デフォルト: 50）
            include_tables: 表も検索対象に含めるか（デフォルト: true）
    """
    file_path = args.get("file_path", "")
    query = args.get("query", "")
    case_sensitive = args.get("case_sensitive", False)
    max_hits = args.get("max_hits", 50)
    include_tables = args.get("include_tables", True)

    old_format_error = _check_old_format(file_path)
    if old_format_error:
        return old_format_error

    if not query:
        return {
            "content": [{"type": "text", "text": "エラー: query（検索キーワード）を指定してください。"}],
            "is_error": True,
        }

    try:
        from docx import Document  # noqa: F401
    except ImportError:
        return {
            "content": [{"type": "text", "text": "エラー: python-docxライブラリがインストールされていません。"}],
            "is_error": True,
        }

    try:
        content, filename, _ = await workspace_service.download_file(
            tenant_id, conversation_id, file_path
        )

        result = search_document(
            content,
            query,
            case_sensitive=case_sensitive,
            max_hits=max_hits,
            include_tables=include_tables,
        )

        result_lines = [
            f"# 検索結果: \"{result['query']}\"",
            f"ヒット数: {result['total_hits']}",
            "",
        ]

        if result['hits']:
            for hit in result['hits']:
                if hit['location_type'] == 'heading':
                    result_lines.append(f"## 見出し (段落{hit['para_index']}, レベル{hit['heading_level']})")
                elif hit['location_type'] == 'paragraph':
                    result_lines.append(f"## 段落 {hit['para_index']}")
                else:
                    result_lines.append(f"## 表{hit['table_index']}")

                result_lines.append(f"コンテキスト: {hit['context']}")
                result_lines.append("")
        else:
            result_lines.append("検索結果はありませんでした。")

        return {
            "content": [{"type": "text", "text": "\n".join(result_lines)}],
        }
    except FileNotFoundError:
        return {
            "content": [{"type": "text", "text": f"ファイルが見つかりません: {file_path}"}],
            "is_error": True,
        }
    except ValueError as e:
        return {
            "content": [{"type": "text", "text": str(e)}],
            "is_error": True,
        }
    except Exception as e:
        logger.error("Word検索エラー", error=str(e), file_path=file_path, query=query)
        return {
            "content": [{"type": "text", "text": f"検索エラー: {str(e)}"}],
            "is_error": True,
        }
